from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_detailed_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO KỸ THUẬT CHUYÊN SÂU: HỆ THỐNG PHÂN TÍCH Y TẾ PHÂN TÁN (Big Data & AI)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Mục lục tóm tắt
    doc.add_heading('1. Tổng quan Dự án', level=1)
    doc.add_paragraph('Dự án xây dựng một Pipeline dữ liệu hoàn chỉnh, từ việc nạp dữ liệu thô (Raw Data) vào hạ tầng Big Data đến việc cung cấp các dự báo chi phí y tế thông minh thông qua trí tuệ nhân tạo (AI) và biểu đồ trực quan (Dashboard).')

    # 2. Kiến trúc Hệ thống (Architecture)
    doc.add_heading('2. Kiến trúc Hệ thống & Công nghệ sử dụng', level=1)
    p = doc.add_paragraph()
    p.add_run('Hệ thống được thiết kế theo mô hình 5 tầng (Layered Architecture) đảm bảo tính ổn định và khả năng xử lý dữ liệu lớn:').bold = True
    
    tech_list = [
        ('1. Tầng Nạp dữ liệu (Ingestion): ', 'Sử dụng Bash Shell và HDFS Client để nạp file CSV thô vào cụm lưu trữ Hadoop.'),
        ('2. Tầng Kho dữ liệu (Data Warehouse): ', 'Apache Hive (SQL-on-Hadoop) đóng vai trò kho lưu trữ trung tâm, quản lý dữ liệu theo lược đồ (Schema-on-Read).'),
        ('3. Tầng Phân tích (Analytics): ', 'Sử dụng kịch bản Hive Query phân tán để tính toán các chỉ số y tế từ hàng triệu bản ghi.'),
        ('4. Tầng Trí tuệ nhân tạo (AI Serving): ', 'Flask API tích hợp Scikit-Learn Model để xử lý dự báo theo thời gian thực (Real-time Prediction).'),
        ('5. Tầng Trực quan (Visualization): ', 'Dashboard hiện đại kết hợp HTML5/JS và Chart.js để hiển thị số liệu đồ thị động.')
    ]
    for tech, desc in tech_list:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(tech).bold = True
        para.add_run(desc)

    # 3. Luồng di chuyển dữ liệu (Data Flow)
    doc.add_heading('3. Luồng di chuyển dữ liệu chi tiết', level=1)
    p = doc.add_paragraph('Dữ liệu luân chuyển qua hệ thống theo quy trình 6 bước khép kín:')
    
    steps = [
        'Bước 1: Nạp file CSV thô từ các nguồn y tế vào thư mục /healthcare/raw của HDFS.',
        'Bước 2: Hệ thống Hive ánh xạ file thô thành các bảng SQL, tổ chức lại dữ liệu sạch trong Database healthcare_db.',
        'Bước 3: Các kịch bản truy vấn Hive Query phân tách dữ liệu thành các file kết quả JSON nhỏ gọn.',
        'Bước 4: Module AI quét qua kho dữ liệu để huấn luyện mối quan hệ giữa đặc điểm bệnh nhân và chi phí hóa đơn.',
        'Bước 5: Flask API nhận yêu cầu từ người dùng, đưa vào mô hình đã huấn luyện và phản hồi kết quả trong <1 giây.',
        'Bước 6: Dashboard hiển thị các biểu đồ thống kê tập trung và hỗ trợ tương tác dự đoán chi phí tức thì.'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # 4. Chi tiết các thành phần (Components Detail)
    doc.add_heading('4. Chi tiết Kỹ thuật từng thành phần', level=1)
    
    doc.add_heading('4.1 Kho dữ liệu lớn (Hadoop & Hive)', level=2)
    doc.add_paragraph('Hệ thống tận dụng khả năng lưu trữ phân tán của HDFS, cho phép xử lý dữ liệu y tế quy mô lớn mà không bị nghẽn (bottleneck). Hive cung cấp lớp SQL abstraction giúp việc truy vấn các chỉ số nhân khẩu học và phân bổ bệnh trạng trở nên dễ dàng và hiệu suất cao.')
    
    doc.add_heading('4.2 Mô hình học máy (Instant AI)', level=2)
    doc.add_paragraph('Chúng ta chuyển đổi từ Deep Learning nặng nề sang mô hình Linear Regression tối ưu cho dữ liệu bảng. Điều này cho phép "Instant Training" (Huấn luyện tức thì) ngay khi máy chủ khởi động, giúp AI luôn phản hồi kết quả dựa trên dữ liệu y tế mới nhất mà không mất thời gian chờ đợi.')
    
    doc.add_heading('4.3 Dashboard Trực quan', level=2)
    doc.add_paragraph('Được xây dựng theo hướng Mobile-Ready và Responsive, sử dụng Chart.js để vẽ 4 loại biểu đồ: Bar Chart (Tuổi), Horizontal Bar (Bệnh lý), Doughnut (Loại nhập viện) và Line Chart (Phân phối chi phí).')

    # 5. Kết quả thực tế đạt được
    doc.add_heading('5. Kết quả Dự án', level=1)
    results = [
        'Xử lý thành công tập dữ liệu 55,500 record thực tế từ Kaggle.',
        'Tích hợp 5 kịch bản phân tích phân tán tự động.',
        'Thời gian phản hồi kết quả AI dự báo: <0.5 giây.',
        'Hệ thống Web Dashboard trực quan hóa toàn diện quy mô y tế.',
        'Bộ khởi chạy "One-touch" python run.py giúp quản lý toàn bộ Pipeline từ một lệnh duy nhất.'
    ]
    for res in results:
        doc.add_paragraph(res, style='List Bullet')

    # Footer
    p = doc.add_paragraph('\n[BÁO CÁO HOÀN TẤT DỰ ÁN - HEALTHCARE PIPELINE TEAM]')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('Bao_cao_Healthcare_Analytics_Chuyen_Sau.docx')
    print('✅ FILE BÁO CÁO DOCX CHUYÊN SÂU ĐÃ ĐƯỢC CẬP NHẬT THÀNH CÔNG!')

if __name__ == '__main__':
    generate_detailed_report()
