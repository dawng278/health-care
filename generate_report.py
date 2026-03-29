from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
title = doc.add_heading('BÁO CÁO DỰ ÁN: PIPELINE PHÂN TÍCH VÀ DỰ BÁO CHI PHÍ Y TẾ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

content = [
    ('1. Đầu vào & Đầu ra (Input/Output)', '- Đầu vào: healthcare_dataset.csv (55,500 records từ Kaggle).\n- Đầu ra: Dashboard phân tích Big Data và AI Prediction.'),
    ('2. Mục đích (Purpose)', 'Xây dựng một nền tảng phân tích chi phí y tế toàn diện dựa trên kiến trúc Big Data (Hadoop/Hive) và AI (Machine Learning).'),
    ('3. Quy trình thu thập dữ liệu (Data Ingestion)', 'Sử dụng kịch bản HDFS (Bash shell) để nạp dữ liệu thô vào hạ tầng phân tán một cách đáng tin cậy.'),
    ('4. Định dạng lưu trữ (Storage Formats)', 'Dữ liệu được quản lý đa tầng: Raw (CSV trên HDFS), Warehouse (Bảng SQL trong Hive), và Presentation (JSON cho Web).'),
    ('5. Làm sạch và tiền xử lý (Cleaning)', 'Sử dụng Hive Query để xử lý định dạng thời gian và Python Scikit-learn để chuẩn hóa dữ liệu mô hình AI.'),
    ('6. Datapack & Quản lý (Architecture)', 'Hệ thống dùng Hadoop HDFS để lưu trữ quy mô lớn và Hive để quản lý cấu trúc dữ liệu theo lược đồ (Data Warehouse).'),
    ('7. Huấn luyện Machine Learning (Training)', 'Huấn luyện mạng Neural dựa trên Scikit-learn (Linear Regression) cho phép phản hồi kết quả dự báo thời hạn ngay lập tức (<1s).'),
    ('8. Data Warehouse (DW)', 'Tích hợp dữ liệu sạch vào Database "healthcare_db" trong Hive, đóng vai trò là "Single Source of Truth" cho phân tích.'),
    ('9. Trực quan hóa dữ liệu (Visualization)', 'Dashboard hiện đại xây dựng trên HTML/JS và Chart.js, cung cấp 4 biểu đồ tương tác thời gian thực.'),
    ('10. Kết quả đạt được (Results)', 'Pipeline vận hành hoàn chỉnh trên 55.5k records. Đạt được khả năng AI dự báo hóa đơn ngay lập tức với giao diện chuyên nghiệp.')
]

for heading, body in content:
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)

doc.save('Bao_cao_Healthcare_Analytics.docx')
print('✅ FILE DOCX ĐÃ ĐƯỢC TẠO THÀNH CÔNG!')
